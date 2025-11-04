import streamlit as st
import re
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import os
import tempfile
from io import BytesIO
import base64

# 设置页面配置
st.set_page_config(
    page_title="图号匹配工具",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)


class FigureNumberProcessor:
    def __init__(self):
        self.found_figures = []
        self.matched_figures = []
        self.unmatched_figures = []
        self.bom_list = []

    def extract_figure_numbers_from_docx(self, docx_content):
        """
        从Word文档中提取图号，只保留包含"-"的英文数字编号
        """
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(docx_content)
                tmp_file_path = tmp_file.name

            doc = Document(tmp_file_path)
            figures = []

            st.info("正在从Word文档中提取图号...")

            # 定义匹配模式 - 只匹配包含"-"的图号
            patterns = [
                r'\b[A-Z]{1,4}\d{1,3}-[A-Z0-9]{4,10}(?:-[A-Z0-9]{1,3})?\b',
                r'\b[A-Z]?\d{5,6}-[A-Z0-9]{6,7}\b',
                r'\b[A-Z0-9]{4,8}-[A-Z0-9]{3,8}\b',
            ]

            # 提取段落中的图号
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                if text.strip():
                    for pattern in patterns:
                        matches = re.finditer(pattern, text)
                        for match in matches:
                            figure = match.group()
                            if (any(c.isalpha() for c in figure) and
                                    any(c.isdigit() for c in figure) and
                                    '-' in figure):
                                figures.append({
                                    'text': figure,
                                    'type': 'paragraph',
                                    'para_idx': para_idx,
                                    'position': match.start(),
                                    'original_text': text
                                })

            # 提取表格中的图号
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            text = paragraph.text
                            if text.strip():
                                for pattern in patterns:
                                    matches = re.finditer(pattern, text)
                                    for match in matches:
                                        figure = match.group()
                                        if (any(c.isalpha() for c in figure) and
                                                any(c.isdigit() for c in figure) and
                                                '-' in figure):
                                            figures.append({
                                                'text': figure,
                                                'type': 'table',
                                                'table_idx': table_idx,
                                                'row_idx': row_idx,
                                                'cell_idx': cell_idx,
                                                'para_idx': para_idx,
                                                'position': match.start(),
                                                'original_text': text
                                            })

            # 清理临时文件
            os.unlink(tmp_file_path)

            # 去重
            unique_figures = []
            seen = set()
            for fig in figures:
                if fig['text'] not in seen:
                    seen.add(fig['text'])
                    unique_figures.append(fig)

            self.found_figures = unique_figures
            st.success(f"共提取到 {len(unique_figures)} 个唯一图号")
            return unique_figures

        except Exception as e:
            st.error(f"读取Word文档时出错: {e}")
            return []

    def load_bom_from_excel(self, excel_content, bom_column='BOM', sheet_name=0):
        """
        从Excel文件加载BOM列表
        """
        try:
            st.info("正在从Excel文件加载BOM数据...")

            # 读取Excel文件
            df = pd.read_excel(BytesIO(excel_content), sheet_name=sheet_name, engine='openpyxl')

            # 查找BOM列
            bom_data = []
            if bom_column in df.columns:
                bom_series = df[bom_column].dropna()
                bom_data = bom_series.astype(str).tolist()
            else:
                # 如果没有找到指定列，尝试查找包含'BOM'的列
                found_column = None
                for col in df.columns:
                    if 'bom' in str(col).lower():
                        found_column = col
                        break

                if found_column:
                    bom_series = df[found_column].dropna()
                    bom_data = bom_series.astype(str).tolist()
                    st.info(f"自动找到BOM列: {found_column}")
                else:
                    st.error("未找到BOM列，请检查Excel文件结构")
                    return []

            # 清理数据
            cleaned_bom_data = []
            for item in bom_data:
                item = str(item).strip()
                if item and item != 'nan' and item != 'None':
                    cleaned_bom_data.append(item)

            self.bom_list = list(set(cleaned_bom_data))
            st.success(f"从Excel文件加载了 {len(self.bom_list)} 个BOM编号")
            return self.bom_list

        except Exception as e:
            st.error(f"读取Excel文件时出错: {e}")
            return []

    def match_figures_with_bom(self):
        """
        将提取的图号与BOM列表进行匹配（子串匹配）
        """
        if not self.found_figures or not self.bom_list:
            st.error("错误：请先加载Word文档和Excel文件")
            return

        st.info("开始匹配图号与BOM列表（子串匹配）...")
        self.matched_figures = []
        self.unmatched_figures = []

        for fig_info in self.found_figures:
            figure = fig_info['text']
            matched = False

            # 子串匹配：检查图号是否是BOM列表中某个字符串的子串
            for bom_item in self.bom_list:
                if figure in bom_item:
                    matched = True
                    break

            if matched:
                self.matched_figures.append(fig_info)
            else:
                self.unmatched_figures.append(fig_info)

        st.success(
            f"匹配完成！匹配的图号: {len(self.matched_figures)} 个，不匹配的图号: {len(self.unmatched_figures)} 个")

    def mark_unmatched_in_word(self, original_content):
        """
        在Word文档中标记不匹配的图号（紫色背景）
        """
        try:
            st.info("正在标记不匹配的图号...")

            # 创建临时输入文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_input:
                tmp_input.write(original_content)
                tmp_input_path = tmp_input.name

            doc = Document(tmp_input_path)

            # 标记段落中的不匹配图号
            for fig_info in self.unmatched_figures:
                if fig_info['type'] == 'paragraph':
                    para_idx = fig_info['para_idx']
                    figure = fig_info['text']

                    if para_idx < len(doc.paragraphs):
                        paragraph = doc.paragraphs[para_idx]
                        original_text = paragraph.text

                        # 清空段落并重新添加带格式的文本
                        paragraph.clear()

                        # 找到图号位置并标记
                        start_pos = original_text.find(figure)
                        if start_pos != -1:
                            # 添加图号前的文本
                            if start_pos > 0:
                                paragraph.add_run(original_text[:start_pos])

                            # 添加标记的图号（紫色背景）
                            run = paragraph.add_run(figure)
                            run.font.color.rgb = RGBColor(255, 255, 255)  # 白色文字
                            run.font.highlight_color = WD_COLOR_INDEX.GRAY_50  # 使用灰色作为近似紫色

                            # 添加图号后的文本
                            if start_pos + len(figure) < len(original_text):
                                paragraph.add_run(original_text[start_pos + len(figure):])

            # 标记表格中的不匹配图号
            for fig_info in self.unmatched_figures:
                if fig_info['type'] == 'table':
                    table_idx = fig_info['table_idx']
                    row_idx = fig_info['row_idx']
                    cell_idx = fig_info['cell_idx']
                    para_idx = fig_info['para_idx']
                    figure = fig_info['text']

                    if (table_idx < len(doc.tables) and
                            row_idx < len(doc.tables[table_idx].rows) and
                            cell_idx < len(doc.tables[table_idx].rows[row_idx].cells)):

                        cell = doc.tables[table_idx].rows[row_idx].cells[cell_idx]
                        if para_idx < len(cell.paragraphs):
                            paragraph = cell.paragraphs[para_idx]
                            original_text = paragraph.text

                            # 清空段落并重新添加带格式的文本
                            paragraph.clear()

                            # 找到图号位置并标记
                            start_pos = original_text.find(figure)
                            if start_pos != -1:
                                # 添加图号前的文本
                                if start_pos > 0:
                                    paragraph.add_run(original_text[:start_pos])

                                # 添加标记的图号（紫色背景）
                                run = paragraph.add_run(figure)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.highlight_color = WD_COLOR_INDEX.GRAY_50

                                # 添加图号后的文本
                                if start_pos + len(figure) < len(original_text):
                                    paragraph.add_run(original_text[start_pos + len(figure):])

            # 保存到临时输出文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_output:
                doc.save(tmp_output.name)
                with open(tmp_output.name, 'rb') as f:
                    marked_content = f.read()

            # 清理临时文件
            os.unlink(tmp_input_path)
            os.unlink(tmp_output.name)

            st.success("标记完成！")
            return marked_content

        except Exception as e:
            st.error(f"标记Word文档时出错: {e}")
            return None


def main():
    # 页面标题
    st.title("🔍 图号匹配工具")
    st.markdown("---")

    # 初始化session state
    if 'processor' not in st.session_state:
        st.session_state.processor = FigureNumberProcessor()
    if 'processed' not in st.session_state:
        st.session_state.processed = False

    # 侧边栏 - 使用说明
    with st.sidebar:
        st.header("使用说明")
        st.markdown("""
        1. **上传Word文档** - 包含图号的文档
        2. **上传Excel文件** - 包含BOM数据的xlsm文件  
        3. **配置参数** - 选择BOM列和工作表
        4. **开始处理** - 自动匹配并标记

        **图号格式示例:**
        - Y24-1205100-19
        - S01200-3900750  
        - 901000-39005R1
        - YCFP-011
        """)

        st.markdown("---")
        st.header("匹配规则")
        st.markdown("""
        - 只提取包含`-`的图号
        - 图号必须是BOM项的子串
        - 不匹配的图号会用紫色背景标记
        """)

    # 文件上传区域
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📄 上传Word文档")
        word_file = st.file_uploader("选择Word文档", type=['docx'], key="word")

    with col2:
        st.subheader("📊 上传Excel文件")
        excel_file = st.file_uploader("选择Excel文件", type=['xlsm', 'xlsx'], key="excel")

    # 配置参数
    st.subheader("⚙️ 配置参数")
    config_col1, config_col2 = st.columns(2)

    with config_col1:
        bom_column = st.text_input("BOM列名称", value="BOM", help="Excel文件中BOM数据所在的列名")

    with config_col2:
        sheet_name = st.text_input("工作表名称", value="0", help="工作表名称或序号（0表示第一个工作表）")

    # 处理按钮
    if st.button("🚀 开始处理", type="primary", use_container_width=True):
        if not word_file or not excel_file:
            st.error("请先上传Word文档和Excel文件！")
            return

        processor = st.session_state.processor

        # 处理Word文档
        with st.spinner("正在处理Word文档..."):
            word_content = word_file.getvalue()
            figures = processor.extract_figure_numbers_from_docx(word_content)

        if not figures:
            st.error("未找到符合要求的图号！")
            return

        # 处理Excel文件
        with st.spinner("正在处理Excel文件..."):
            excel_content = excel_file.getvalue()
            try:
                sheet = int(sheet_name) if sheet_name.isdigit() else sheet_name
            except:
                sheet = 0

            bom_list = processor.load_bom_from_excel(excel_content, bom_column, sheet)

        if not bom_list:
            st.error("未找到BOM数据！")
            return

        # 匹配图号
        with st.spinner("正在匹配图号..."):
            processor.match_figures_with_bom()

        st.session_state.processed = True

    # 显示结果
    if st.session_state.processed:
        processor = st.session_state.processor

        st.markdown("---")
        st.subheader("📊 处理结果")

        # 统计信息
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("提取图号总数", len(processor.found_figures))
        with col2:
            st.metric("匹配图号数量", len(processor.matched_figures),
                      delta=f"{len(processor.matched_figures) / len(processor.found_figures) * 100:.1f}%" if processor.found_figures else 0)
        with col3:
            st.metric("不匹配图号数量", len(processor.unmatched_figures),
                      delta=f"-{len(processor.unmatched_figures)}", delta_color="inverse")

        # 结果显示
        tab1, tab2, tab3 = st.tabs(["📋 所有图号", "✅ 匹配图号", "❌ 不匹配图号"])

        with tab1:
            if processor.found_figures:
                all_figures_df = pd.DataFrame([
                    {'图号': fig['text'], '类型': fig['type'], '位置': f"段落{fig['para_idx'] + 1}" if fig[
                                                                                                           'type'] == 'paragraph' else f"表格{fig['table_idx'] + 1}"}
                    for fig in processor.found_figures
                ])
                st.dataframe(all_figures_df, use_container_width=True)
            else:
                st.info("未找到图号")

        with tab2:
            if processor.matched_figures:
                matched_figures_df = pd.DataFrame([
                    {'图号': fig['text'], '类型': fig['type'], '位置': f"段落{fig['para_idx'] + 1}" if fig[
                                                                                                           'type'] == 'paragraph' else f"表格{fig['table_idx'] + 1}"}
                    for fig in processor.matched_figures
                ])
                st.dataframe(matched_figures_df, use_container_width=True)
            else:
                st.info("没有匹配的图号")

        with tab3:
            if processor.unmatched_figures:
                unmatched_figures_df = pd.DataFrame([
                    {'图号': fig['text'], '类型': fig['type'], '位置': f"段落{fig['para_idx'] + 1}" if fig[
                                                                                                           'type'] == 'paragraph' else f"表格{fig['table_idx'] + 1}"}
                    for fig in processor.unmatched_figures
                ])
                st.dataframe(unmatched_figures_df, use_container_width=True)
            else:
                st.success("所有图号都匹配成功！")

        # 下载标记后的文档
        if processor.unmatched_figures:
            st.markdown("---")
            st.subheader("📥 下载结果")

            with st.spinner("正在生成标记文档..."):
                marked_content = processor.mark_unmatched_in_word(word_content)

                if marked_content:
                    # 提供下载链接
                    st.download_button(
                        label="📄 下载标记后的Word文档",
                        data=marked_content,
                        file_name=f"标记版_{word_file.name}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

        # 生成报告
        st.markdown("---")
        st.subheader("📈 详细报告")

        if processor.found_figures:
            report_data = []
            for fig in processor.found_figures:
                status = "匹配" if fig in processor.matched_figures else "不匹配"
                location = f"段落{fig['para_idx'] + 1}" if fig['type'] == 'paragraph' else f"表格{fig['table_idx'] + 1}"

                report_data.append({
                    '图号': fig['text'],
                    '状态': status,
                    '位置': location,
                    '类型': fig['type']
                })

            report_df = pd.DataFrame(report_data)

            # 提供Excel报告下载
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                report_df.to_excel(writer, sheet_name='匹配报告', index=False)

                # 添加统计表
                stats_df = pd.DataFrame({
                    '统计项': ['提取图号总数', '匹配图号数', '不匹配图号数', '匹配率'],
                    '数值': [
                        len(processor.found_figures),
                        len(processor.matched_figures),
                        len(processor.unmatched_figures),
                        f"{len(processor.matched_figures) / len(processor.found_figures) * 100:.1f}%" if processor.found_figures else "0%"
                    ]
                })
                stats_df.to_excel(writer, sheet_name='统计信息', index=False)

            excel_buffer.seek(0)

            st.download_button(
                label="📊 下载Excel详细报告",
                data=excel_buffer,
                file_name=f"图号匹配报告_{word_file.name.split('.')[0]}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )


if __name__ == "__main__":
    main()